import os
from docx.text.paragraph import Paragraph
from docx.document import Document
from docx.table import _Cell, Table
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
import docx
import random


def iter_block_items(parent):
    if isinstance(parent, Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


# ========================================================
# Get docs directory
# ========================================================

path = "docs"
dir_list = os.listdir(path)
if len(dir_list) == 0:
    print("No files found in docs directory")
    exit()
print(dir_list)

# ========================================================
# Select Docx Files
# ========================================================

for file in dir_list:
    filename = file[:-5]
    if file.endswith(".docx"):
        doc = docx.Document(f"docs/{file}")

        for elem in doc.inline_shapes:
            print(elem)
        print(f"Processing {len(doc.inline_shapes)}")
        syllabus = docx.Document()
        code = ""
        name = "General"
        element_list = []
        skip = False
        list_num = 1
        for block in iter_block_items(doc):
            element_list.append(block)

        for i in range(len(element_list)):
            block = element_list[i]
            print(i)

            try:
                if isinstance(block, Paragraph):
                    if block.text.strip() == "":
                        continue
                    inserted_p = syllabus._body._body._insert_p(block._p)
                    if block._p.get_or_add_pPr().numPr:
                        inserted_p.style = "ListNumber"
                elif isinstance(block, Table):
                    if skip:
                        skip = False
                        continue
                    if (
                        "Mapping between COs".lower()
                        in block.rows[0].cells[0].text.lower()
                    ):
                        print(block, code, name)
                        avg = []
                        avg_count = []
                        pso_avg = []
                        for z in range(0, 12):
                            avg.append(0)
                            avg_count.append(0)

                        for z in range(0, 3):
                            pso_avg.append(0)
                        course_code_number = code[3:]
                        cos = len(block.rows)

                        mapping_table = syllabus.add_table(rows=cos, cols=18)
                        mapping_table.style = 'Table Grid'
                        mapping_table.autofit = True
                        header = [
                            "Course Code",
                            "Course Name",
                            "COs",
                            "PO1",
                            "PO2",
                            "PO3",
                            "PO4",
                            "PO5",
                            "PO6",
                            "PO7",
                            "PO8",
                            "PO9",
                            "PO10",
                            "PO11",
                            "PO12",
                            "PSO1",
                            "PSO2",
                            "PSO3",
                        ]
                        for j in range(len(header)):
                            mapping_table.rows[0].cells[j].text = header[j]

                        for j in range(2, cos):
                            o_j = j - 1
                            data_row = block.rows[j]
                            output_row = mapping_table.rows[o_j]
                            if j - 1 == 1:
                                mapping_table.rows[1].cells[0].text = code
                                mapping_table.rows[1].cells[1].text = name

                            else:
                                mapping_table.rows[o_j].cells[0].text = ""
                                mapping_table.rows[o_j].cells[1].text = ""

                            mapping_table.rows[o_j].cells[
                                2
                            ].text = f"CO{course_code_number}.{j-1}"
                            origin_pos = data_row.cells[2].text.split(",")
                            origin_pos = [_.strip() for _ in origin_pos]
                            pos = origin_pos + [
                                "PO1",
                                "PO2",
                                "PO3",
                                "PO4",
                                "PO5",
                                "PO6",
                                "PO7",
                                "PO9",
                                "PO12",
                            ]
                            pos = [_.strip() for _ in pos]
                            pos = list(set(pos))
                            for k in range(0, 12):
                                po = f"PO{k+1}"
                                if po in pos:
                                    if po in origin_pos:
                                        val = random.randint(2, 3)
                                    else:
                                        val = random.randint(1, 3)

                                    avg[k] = avg[k] + val
                                    avg_count[k] = avg_count[k] + 1
                                else:
                                    val = "-"
                                output_row.cells[k + 3].text = str(val)

                            for k in range(0, 3):
                                val = random.randint(1, 3)
                                pso_avg[k] = pso_avg[k] + val
                                output_row.cells[k + 15].text = str(val)
                            skip = True

                        avg_new = []

                        for t in range(len(avg)):
                            val = 0
                            if avg_count[t] > 0:
                                val = round(avg[t] / avg_count[t], 2)
                            avg_new.append(val)
                        pso_avg_new = [round(x / (cos - 2), 2) for x in pso_avg]

                        mapping_table.rows[cos - 1].cells[
                            2
                        ].text = f"CO{course_code_number}"
                        for k in range(0, 12):
                            if avg_new[k] > 0:
                                mapping_table.rows[cos - 1].cells[k + 3].text = str(
                                    avg_new[k]
                                )
                            else:
                                mapping_table.rows[cos - 1].cells[k + 3].text = "-"

                        for k in range(0, 3):
                            if pso_avg_new[k] > 0:
                                mapping_table.rows[cos - 1].cells[k + 15].text = str(
                                    pso_avg_new[k]
                                )
                            else:
                                mapping_table.rows[cos - 1].cells[k + 15].text = "-"
                        mapping_table.cell(1, 0).merge(mapping_table.cell(cos - 1, 0))
                        mapping_table.cell(1, 1).merge(mapping_table.cell(cos - 1, 1))

                    else:
                        cells = block.rows[0].cells
                        if len(cells) == 6:
                            code = "".join(
                                i
                                for i in cells[0].text.strip()
                                if i not in ["/", ":", "\\", "Course code"]
                            )
                            name = "".join(
                                i
                                for i in cells[1].text.strip()
                                if i not in ["/", ":", "\\", "Course Name"]
                            )
                            syllabus.add_page_break()

                        paragraph = syllabus.add_paragraph()
                        paragraph._p.addnext(block._tbl)
            except Exception as e:
                print(e)
        path = f"output/{filename}"
        if not os.path.exists(path):
            os.makedirs(path)
        syllabus.save(f"output/{filename}.docx")
    else:
        print("Not docx file")
